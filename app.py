import os, uuid, io
import logging
import tempfile
import base64
import traceback
import csv
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware

# Configure logging with more detail
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Log startup
logger.info("=" * 50)
logger.info("Starting Document to Markdown converter application...")
logger.info("=" * 50)

# Try to import required libraries with error handling
logger.info("Importing required libraries...")

# Global variables for optional dependencies
Image = None
pd = None
openpyxl = None
Document = None
markdownify = None

try:
    from PIL import Image
    logger.info("✅ Pillow imported successfully")
except ImportError as e:
    logger.warning(f"⚠️ Failed to import Pillow: {e}")
    logger.warning("Image processing functionality will be disabled")

try:
    import pandas as pd
    logger.info("✅ pandas imported successfully")
except ImportError as e:
    logger.warning(f"⚠️ Failed to import pandas: {e}")
    logger.warning("CSV/Excel functionality will be disabled")

try:
    import openpyxl
    logger.info("✅ openpyxl imported successfully")
except ImportError as e:
    logger.warning(f"⚠️ Failed to import openpyxl: {e}")
    logger.warning("Excel functionality will be disabled")

try:
    from docx import Document
    logger.info("✅ python-docx imported successfully")
except ImportError as e:
    logger.warning(f"⚠️ Failed to import python-docx: {e}")
    logger.warning("Word document functionality will be disabled")

try:
    from markdownify import markdownify
    logger.info("✅ markdownify imported successfully")
except ImportError as e:
    logger.warning(f"⚠️ Failed to import markdownify: {e}")
    logger.warning("HTML to markdown conversion will be disabled")

try:
    import openai
except ImportError:
    openai = None
    logger.warning("⚠️ OpenAI library not installed. Summarization will be disabled.")

async def summarize_markdown(text: str, max_chars: int = 3000) -> str:
    """Summarize markdown text to fit within max_chars using OpenAI (new API >=1.0.0)."""
    if openai is None:
        logger.error("OpenAI library not available for summarization.")
        raise RuntimeError("OpenAI library not available.")
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        logger.error("OPENAI_API_KEY not set in environment.")
        raise RuntimeError("OPENAI_API_KEY not set.")
    prompt = (
        f"Summarize the following markdown content in at most {max_chars} characters. "
        "Preserve as much structure and key information as possible.\n\n" + text
    )
    import asyncio
    try:
        logger.info(f"Calling OpenAI (new API) to summarize markdown (original length: {len(text)} chars)...")
        client = openai.OpenAI(api_key=api_key)
        response = await asyncio.to_thread(
            client.chat.completions.create,
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=16384,  # <-- updated to max allowed
            temperature=0,
        )
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            summary = response.choices[0].message.content.strip()
        else:
            logger.error("OpenAI response missing expected content for summarization.")
            raise RuntimeError("OpenAI response missing expected content for summarization.")
        logger.info(f"✅ OpenAI summarization complete (summary length: {len(summary)} chars)")
        return summary
    except Exception as e:
        logger.error(f"❌ OpenAI summarization failed: {e}")
        raise

logger.info("Library import check completed!")

app = FastAPI(title="Document to Markdown Converter")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variables


def image_to_base64(image) -> str:
    """Convert PIL Image to base64 string"""
    if Image is None:
        raise ImportError("Pillow (PIL) is not available")
    
    try:
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        base64_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
        logger.info(f"✅ Image converted to base64 ({len(base64_str)} characters)")
        return base64_str
    except Exception as e:
        logger.error(f"❌ Error converting image to base64: {e}")
        raise

def docx_to_markdown(docx_data: bytes) -> str:
    """Convert DOCX data to markdown"""
    if Document is None:
        raise ImportError("python-docx is not available")
    
    temp_file_path = None
    try:
        logger.info("Converting DOCX to markdown...")
        
        # Save DOCX data to temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_file.write(docx_data)
            temp_file_path = temp_file.name
        
        # Open DOCX document
        doc = Document(temp_file_path)
        
        # Extract text and formatting
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                # Check for heading styles
                if para.style and getattr(para.style, 'name', None) and isinstance(para.style.name, str) and para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                    paragraphs.append(f"{'#' * int(level)} {para.text}")
                else:
                    paragraphs.append(para.text)
        
        # Extract table data
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                # Convert table to markdown
                if len(table_data) > 0:
                    headers = table_data[0]
                    body = table_data[1:]
                    
                    # Create markdown table
                    table_lines = []
                    table_lines.append('| ' + ' | '.join(headers) + ' |')
                    table_lines.append('|' + '|'.join(['---'] * len(headers)) + '|')
                    
                    for row in body:
                        # Pad row if shorter than headers
                        while len(row) < len(headers):
                            row.append('')
                        # Truncate if longer
                        row = row[:len(headers)]
                        table_lines.append('| ' + ' | '.join(row) + ' |')
                    
                    tables.append('\n'.join(table_lines))
        
        # Combine paragraphs and tables
        content = '\n\n'.join(paragraphs)
        if tables:
            content += '\n\n' + '\n\n'.join(tables)
        
        logger.info(f"✅ DOCX converted to markdown ({len(content)} characters)")
        return content
        
    except Exception as e:
        logger.error(f"❌ Error converting DOCX to markdown: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise
    finally:
        if temp_file_path:
            try:
                os.unlink(temp_file_path)
                logger.info("✅ Temporary DOCX file cleaned up")
            except Exception as e:
                logger.warning(f"⚠️ Failed to clean up temporary DOCX file: {e}")

def csv_to_markdown(csv_data: bytes) -> str:
    """Convert CSV data to markdown table"""
    try:
        logger.info("Converting CSV to markdown...")
        
        # Decode CSV data
        csv_text = csv_data.decode('utf-8')
        
        # Parse CSV
        lines = csv_text.strip().split('\n')
        if not lines:
            raise ValueError("Empty CSV file")
        
        # Parse CSV properly using csv module
        csv_reader = csv.reader(lines)
        rows = list(csv_reader)
        
        if not rows:
            raise ValueError("No data in CSV file")
        
        header = rows[0]
        body = rows[1:]
        
        logger.info(f"✅ CSV parsed: {len(header)} columns, {len(body)} rows")
        
        # Create markdown table
        def make_row(cols):
            return '| ' + ' | '.join(str(col).strip() for col in cols) + ' |'
        
        md_lines = []
        md_lines.append(make_row(header))
        md_lines.append('|' + '|'.join(['---'] * len(header)) + '|')
        
        for row in body:
            # Pad row if it's shorter than header
            while len(row) < len(header):
                row.append('')
            # Truncate row if it's longer than header
            row = row[:len(header)]
            md_lines.append(make_row(row))
        
        result = '\n'.join(md_lines)
        logger.info(f"✅ CSV converted to markdown ({len(result)} characters)")
        return result
        
    except Exception as e:
        logger.error(f"❌ Error converting CSV to markdown: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise

def excel_to_markdown(excel_data: bytes) -> str:
    """Convert Excel data to markdown table"""
    if pd is None:
        raise ImportError("pandas is not available")
    if openpyxl is None:
        raise ImportError("openpyxl is not available")
    
    temp_file_path = None
    try:
        logger.info("Converting Excel to markdown...")
        
        # Save Excel data to temporary file
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp_file:
            temp_file.write(excel_data)
            temp_file_path = temp_file.name
        
        # Read Excel file with pandas
        df = pd.read_excel(temp_file_path, engine='openpyxl')
        
        logger.info(f"✅ Excel parsed: {len(df.columns)} columns, {len(df)} rows")
        
        # Convert to markdown - handle missing tabulate dependency
        try:
            markdown_table = df.to_markdown(index=False)
        except ImportError:
            # Fallback: create markdown table manually
            logger.info("tabulate not available, creating markdown table manually")
            headers = df.columns.tolist()
            rows = df.values.tolist()
            
            # Create markdown table
            table_lines = []
            table_lines.append('| ' + ' | '.join(str(col) for col in headers) + ' |')
            table_lines.append('|' + '|'.join(['---'] * len(headers)) + '|')
            
            for row in rows:
                # Convert all values to strings and escape pipes
                row_str = [str(cell).replace('|', '\\|') for cell in row]
                table_lines.append('| ' + ' | '.join(row_str) + ' |')
            
            markdown_table = '\n'.join(table_lines)
        if markdown_table is None:
            markdown_table = ""
        logger.info(f"✅ Excel converted to markdown ({len(markdown_table)} characters)")
        return markdown_table
        
    except Exception as e:
        logger.error(f"❌ Error converting Excel to markdown: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise
    finally:
        if temp_file_path:
            try:
                os.unlink(temp_file_path)
                logger.info("✅ Temporary Excel file cleaned up")
            except Exception as e:
                logger.warning(f"⚠️ Failed to clean up temporary Excel file: {e}")

@app.on_event("startup")
async def startup_event():
    """Initialize services on startup"""
    logger.info("=" * 50)
    logger.info("Application starting up...")
    logger.info("=" * 50)
    
    try:
        # Remove openai import and related variables
        # Remove initialize_openai and openai_client
        # Remove call_openai_with_text and call_openai_with_image
        # Remove OpenAI logic from /upload and health check
        pass # No OpenAI initialization needed
    except Exception as e:
        logger.error(f"❌ Startup error: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")

@app.get("/", response_class=HTMLResponse)
def index():
    logger.info("Serving index page")
    try:
        with open("static/index.html", encoding="utf8") as f:
            content = f.read()
            logger.info(f"✅ Index page loaded successfully ({len(content)} characters)")
            return content
    except Exception as e:
        logger.error(f"❌ Error reading index.html: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(500, "Internal server error")

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    logger.info("=" * 50)
    logger.info(f"Upload request received for file: {file.filename}")
    logger.info("=" * 50)
    
    try:
        # Read file data
        logger.info("Reading file data...")
        data = await file.read()
        logger.info(f"✅ File read successfully ({len(data)} bytes)")
        
        # Determine file type and process accordingly
        if not file.filename or not file.content_type:
            logger.error("❌ Uploaded file is missing filename or content_type")
            raise HTTPException(400, "Uploaded file is missing filename or content_type")
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        content_type = file.content_type.lower()
        
        logger.info(f"File type: {content_type}, extension: {file_extension}")
        
        if content_type == "application/pdf" or file_extension == "pdf":
            # Process PDF using MarkItDown
            logger.info("Processing PDF file with MarkItDown...")
            try:
                from markitdown import MarkItDown
            except ImportError as e:
                logger.error(f"❌ MarkItDown not installed: {e}")
                raise HTTPException(500, "MarkItDown is not installed on the server.")
            try:
                md_converter = MarkItDown()
                import io
                pdf_stream = io.BytesIO(data)
                result = md_converter.convert(pdf_stream)
                md = result.text_content
                pages_processed = 1  # MarkItDown does not split by page, so treat as one doc
                logger.info(f"✅ PDF converted to markdown using MarkItDown ({len(md)} characters)")
            except Exception as e:
                logger.error(f"❌ Error converting PDF with MarkItDown: {e}")
                logger.error(f"Traceback: {traceback.format_exc()}")
                raise HTTPException(500, f"Failed to convert PDF: {str(e)}")
        
        elif content_type in ["application/vnd.openxmlformats-officedocument.presentationml.presentation", 
                             "application/vnd.ms-powerpoint"] or file_extension in ["pptx", "ppt"]:
            # Process PowerPoint using MarkItDown
            logger.info("Processing PowerPoint file with MarkItDown...")
            try:
                from markitdown import MarkItDown
            except ImportError as e:
                logger.error(f"❌ MarkItDown not installed: {e}")
                raise HTTPException(500, "MarkItDown is not installed on the server.")
            try:
                md_converter = MarkItDown()
                import io
                pptx_stream = io.BytesIO(data)
                result = md_converter.convert(pptx_stream)
                md = result.text_content
                pages_processed = 1  # MarkItDown does not split by slide, so treat as one doc
                logger.info(f"✅ PowerPoint converted to markdown using MarkItDown ({len(md)} characters)")
            except Exception as e:
                logger.error(f"❌ Error converting PowerPoint with MarkItDown: {e}")
                logger.error(f"Traceback: {traceback.format_exc()}")
                raise HTTPException(500, f"Failed to convert PowerPoint: {str(e)}")
        
        elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             "application/msword"] or file_extension in ["docx", "doc"]:
            # Process Word document
            logger.info("Processing Word document...")
            md = docx_to_markdown(data)
            pages_processed = 1
        
        elif content_type == "text/csv" or file_extension == "csv":
            # Process CSV
            logger.info("Processing CSV file...")
            md = csv_to_markdown(data)
            pages_processed = 1
        
        elif content_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             "application/vnd.ms-excel"] or file_extension in ["xlsx", "xls"]:
            # Process Excel
            logger.info("Processing Excel file...")
            md = excel_to_markdown(data)
            pages_processed = 1
        
        else:
            # For other file types, try to extract text and convert
            logger.info("Processing as text file...")
            try:
                text_content = data.decode('utf-8')
                if markdownify is not None:
                    md = markdownify(text_content)
                else:
                    md = text_content
                pages_processed = 1
            except UnicodeDecodeError:
                logger.error(f"❌ Unsupported file type: {content_type}")
                raise HTTPException(400, "Unsupported file type. Please upload PDF, PowerPoint, Word, CSV, Excel, or text files.")
        
        logger.info(f"✅ Conversion successful, markdown length: {len(md)} characters")

        # Summarize if markdown is too long
        summarized = False
        if len(md) > 3000:
            logger.info(f"Markdown exceeds 3000 characters, summarizing...")
            try:
                md = await summarize_markdown(md, max_chars=3000)
                summarized = True
            except Exception as e:
                logger.error(f"Summarization failed, returning original markdown: {e}")
                summarized = False

        # Generate a unique file ID for reference
        file_id = str(uuid.uuid4())
        
        response_data = {
            "file_id": file_id,
            "filename": file.filename,
            "markdown": md,
            "pages_processed": pages_processed,
            "markdown_length": len(md),
            "summarized": summarized
        }
        
        logger.info("✅ Returning successful response")
        return JSONResponse(response_data)
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        logger.info("Re-raising HTTP exception")
        raise
    except Exception as e:
        logger.error(f"❌ Unexpected error during upload: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(500, f"Internal server error: {str(e)}")

# Add health check endpoint
@app.get("/health")
async def health_check():
    logger.info("Health check requested")
    return {
        "status": "healthy"
    }
